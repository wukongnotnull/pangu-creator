"""
素材采集器

支持多种素材格式的读取和解析：
- 文档: PDF, Word (.docx), TXT, Markdown, Excel
- 字幕: SRT, VTT, LRC
- 图片: JPG, PNG, GIF, BMP, WebP, TIFF（OCR 文字识别）
- 音视频: 音频转录（需要 transcribe 模块）
- 网页: URL 抓取
"""

import re
import json
import logging
from pathlib import Path
from typing import List, Optional, Dict, Any, Union
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum

from .models import ContentResult, ContentLanguage

# Module-level logger
logger = logging.getLogger("scripts.search.collector")

def set_log_level(level: int) -> None:
    """Set the logging level for this module."""
    logger.setLevel(level)


def _setup_logging() -> None:
    """Configure module logging with default format."""
    if not logger.handlers:
        handler = logging.StreamHandler()
        formatter = logging.Formatter(
            "%(asctime)s - %(name)s - %(levelname)s - %(message)s"
        )
        handler.setFormatter(formatter)
        logger.addHandler(handler)
        logger.setLevel(logging.INFO)

_setup_logging()


class MaterialType(Enum):
    """素材类型"""
    PDF = "pdf"
    WORD = "word"
    TXT = "txt"
    MARKDOWN = "markdown"
    EXCEL = "excel"
    SUBTITLE = "subtitle"  # srt, vtt, lrc
    AUDIO = "audio"        # mp3, wav, m4a
    VIDEO = "video"        # mp4, mov
    IMAGE = "image"        # jpg, png, webp, etc.
    URL = "url"
    UNKNOWN = "unknown"


@dataclass
class Material:
    """单个素材"""
    path: str                    # 文件路径或URL
    material_type: MaterialType  # 素材类型
    content: str = ""            # 解析后的文本内容
    title: str = ""              # 标题（从文件名或内容提取）
    word_count: int = 0          # 字数
    language: ContentLanguage = ContentLanguage.UNKNOWN
    error: Optional[str] = None
    success: bool = False

    def to_dict(self) -> dict:
        return {
            "path": self.path,
            "type": self.material_type.value,
            "title": self.title,
            "content": self.content,
            "word_count": self.word_count,
            "language": self.language.value,
            "success": self.success,
            "error": self.error,
        }


@dataclass
class CollectionResult:
    """素材采集结果"""
    materials: List[Material] = field(default_factory=list)
    total_files: int = 0
    successful: int = 0
    failed: int = 0
    total_words: int = 0
    completed_at: datetime = field(default_factory=datetime.now)

    def to_dict(self) -> dict:
        return {
            "total_files": self.total_files,
            "successful": self.successful,
            "failed": self.failed,
            "total_words": self.total_words,
            "completed_at": self.completed_at.isoformat(),
            "materials": [m.to_dict() for m in self.materials],
        }


class MaterialCollector:
    """
    素材采集器

    统一接口，支持多种格式：
    - 本地文件：PDF, Word, TXT, Markdown, Excel, 字幕, 音视频
    - 在线内容：URL
    """

    def __init__(self, transcript_enabled: bool = True):
        """
        Args:
            transcript_enabled: 是否启用音视频转录
        """
        self.transcript_enabled = transcript_enabled

    def collect(
        self,
        paths: List[str],
        transcript_enabled: Optional[bool] = None,
    ) -> CollectionResult:
        """
        采集多个素材

        Args:
            paths: 文件路径或URL列表
            transcript_enabled: 是否转录音视频（覆盖构造参数）

        Returns:
            采集结果
        """
        result = CollectionResult()
        result.total_files = len(paths)

        if transcript_enabled is None:
            transcript_enabled = self.transcript_enabled

        for path in paths:
            material = self._collect_single(path, transcript_enabled)
            result.materials.append(material)

            if material.success:
                result.successful += 1
                result.total_words += material.word_count
            else:
                result.failed += 1

        return result

    def _collect_single(
        self,
        path: str,
        transcript_enabled: bool,
    ) -> Material:
        """采集单个素材"""
        material_type = self._detect_type(path)

        material = Material(
            path=path,
            material_type=material_type,
        )

        try:
            if material_type == MaterialType.URL:
                self._collect_url(material)
            elif material_type == MaterialType.PDF:
                self._collect_pdf(material)
            elif material_type == MaterialType.WORD:
                self._collect_word(material)
            elif material_type == MaterialType.TXT:
                self._collect_txt(material)
            elif material_type == MaterialType.MARKDOWN:
                self._collect_markdown(material)
            elif material_type == MaterialType.EXCEL:
                self._collect_excel(material)
            elif material_type == MaterialType.SUBTITLE:
                self._collect_subtitle(material)
            elif material_type == MaterialType.IMAGE:
                self._collect_image(material)
            elif material_type in (MaterialType.AUDIO, MaterialType.VIDEO):
                if transcript_enabled:
                    self._collect_media(material)
                else:
                    material.error = "音视频转录未启用"
            else:
                material.error = f"不支持的格式: {material_type.value}"
        except Exception as e:
            material.error = str(e)

        return material

    def _detect_type(self, path: str) -> MaterialType:
        """检测素材类型"""
        path_lower = path.lower().strip()

        # URL
        if path_lower.startswith(("http://", "https://", "www.")):
            return MaterialType.URL

        # 本地文件
        p = Path(path)
        if not p.exists():
            return MaterialType.UNKNOWN

        ext = p.suffix.lower()

        type_map = {
            ".pdf": MaterialType.PDF,
            ".docx": MaterialType.WORD,
            ".doc": MaterialType.WORD,
            ".txt": MaterialType.TXT,
            ".md": MaterialType.MARKDOWN,
            ".markdown": MaterialType.MARKDOWN,
            ".xlsx": MaterialType.EXCEL,
            ".xls": MaterialType.EXCEL,
            ".csv": MaterialType.EXCEL,
            ".srt": MaterialType.SUBTITLE,
            ".vtt": MaterialType.SUBTITLE,
            ".lrc": MaterialType.SUBTITLE,
            ".mp3": MaterialType.AUDIO,
            ".wav": MaterialType.AUDIO,
            ".m4a": MaterialType.AUDIO,
            ".mp4": MaterialType.VIDEO,
            ".mov": MaterialType.VIDEO,
            ".avi": MaterialType.VIDEO,
            # 图片格式
            ".jpg": MaterialType.IMAGE,
            ".jpeg": MaterialType.IMAGE,
            ".png": MaterialType.IMAGE,
            ".gif": MaterialType.IMAGE,
            ".bmp": MaterialType.IMAGE,
            ".webp": MaterialType.IMAGE,
            ".tiff": MaterialType.IMAGE,
            ".tif": MaterialType.IMAGE,
            ".svg": MaterialType.IMAGE,
        }

        return type_map.get(ext, MaterialType.UNKNOWN)

    def _collect_url(self, material: Material):
        """采集URL内容"""
        from crawl.fetcher import ContentFetcher

        fetcher = ContentFetcher()
        result = fetcher.fetch(material.path)

        material.content = result.content
        material.title = result.title
        material.word_count = result.word_count
        material.language = result.language
        material.success = True

    def _collect_pdf(self, material: Material):
        """采集PDF内容"""
        try:
            import PyPDF2
        except ImportError:
            # 尝试其他库
            try:
                import pypdf
            except ImportError:
                material.error = "需要安装 PyPDF2: pip install PyPDF2"
                return

        path = Path(material.path)
        material.title = path.stem

        try:
            with open(path, "rb") as f:
                if "PyPDF2" in dir():
                    reader = PyPDF2.PdfReader(f)
                else:
                    reader = pypdf.PdfReader(f)

                text_parts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

                content = "\n".join(text_parts)
                material.content = content
                material.word_count = len(content.replace("\n", ""))
                material.language = self._detect_content_lang(content)
                material.success = True
        except Exception as e:
            material.error = f"PDF解析失败: {e}"

    def _collect_word(self, material: Material):
        """采集Word内容"""
        try:
            from docx import Document
        except ImportError:
            material.error = "需要安装 python-docx: pip install python-docx"
            return

        path = Path(material.path)
        material.title = path.stem

        try:
            doc = Document(path)
            paragraphs = [p.text for p in doc.paragraphs]
            content = "\n".join(paragraphs)

            # 也提取表格
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    content += "\n" + " | ".join(cells)

            material.content = content
            material.word_count = len(content.replace("\n", ""))
            material.language = self._detect_content_lang(content)
            material.success = True
        except Exception as e:
            material.error = f"Word解析失败: {e}"

    def _collect_txt(self, material: Material):
        """采集文本文件"""
        path = Path(material.path)
        material.title = path.stem

        try:
            # 尝试不同编码
            for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
                try:
                    content = path.read_text(encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue

            material.content = content
            material.word_count = len(content.replace("\n", ""))
            material.language = self._detect_content_lang(content)
            material.success = True
        except Exception as e:
            material.error = f"文本读取失败: {e}"

    def _collect_markdown(self, material: Material):
        """采集Markdown文件"""
        path = Path(material.path)
        material.title = path.stem

        try:
            content = path.read_text(encoding="utf-8")
            material.content = content
            material.word_count = len(content.replace("\n", ""))
            material.language = self._detect_content_lang(content)
            material.success = True
        except Exception as e:
            material.error = f"Markdown读取失败: {e}"

    def _collect_excel(self, material: Material):
        """采集Excel文件"""
        try:
            import openpyxl
        except ImportError:
            material.error = "需要安装 openpyxl: pip install openpyxl"
            return

        path = Path(material.path)
        material.title = path.stem

        try:
            wb = openpyxl.load_workbook(path)
            parts = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                parts.append(f"## Sheet: {sheet_name}")

                for row in sheet.iter_rows(values_only=True):
                    # 过滤空行
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(cells):
                        parts.append(" | ".join(cells))

            content = "\n".join(parts)
            material.content = content
            material.word_count = len(content.replace("\n", ""))
            material.language = self._detect_content_lang(content)
            material.success = True
        except Exception as e:
            material.error = f"Excel解析失败: {e}"

    def _collect_subtitle(self, material: Material):
        """采集字幕文件"""
        path = Path(material.path)
        material.title = path.stem

        try:
            content = path.read_text(encoding="utf-8")

            # 提取纯文本（去掉时间戳）
            lines = []
            for line in content.split("\n"):
                line = line.strip()
                # 跳过空行、时间戳、序号
                if line and not line.replace(":", "").replace(".", "").replace(",", "").isdigit():
                    if not re.match(r"\d{2}:\d{2}:\d{2}", line):
                        lines.append(line)

            text = " ".join(lines)
            material.content = text
            material.word_count = len(text)
            material.language = self._detect_content_lang(text)
            material.success = True
        except Exception as e:
            material.error = f"字幕解析失败: {e}"

    def _collect_image(self, material: Material):
        """采集图片（OCR识别）"""
        from PIL import Image
        import pytesseract

        path = Path(material.path)
        material.title = path.stem

        try:
            # 打开图片
            img = Image.open(path)

            # OCR 识别
            text = pytesseract.image_to_string(img, lang="chi_sim+eng")

            material.content = text.strip()
            material.word_count = len(text.replace("\n", "").replace(" ", ""))
            material.language = self._detect_content_lang(text)
            material.success = True

        except ImportError as e:
            if "pytesseract" in str(e):
                material.error = "需要安装 OCR 依赖: pip install pytesseract pillow && brew install tesseract"
            elif "PIL" in str(e):
                material.error = "需要安装 Pillow: pip install pillow"
            else:
                material.error = f"缺少依赖: {e}"
        except Exception as e:
            material.error = f"OCR识别失败: {e}"

    def _collect_media(self, material: Material):
        """采集音视频（转录）"""
        from transcribe import AudioTranscriber, YouTubeTranscriber, is_youtube_url

        path = Path(material.path)
        material.title = path.stem

        try:
            transcriber = AudioTranscriber(model="base")

            if is_youtube_url(material.path):
                # YouTube
                yt = YouTubeTranscriber()
                result = yt.get_subtitle(material.path)
                if result.subtitles:
                    content = "\n".join([s["content"] for s in result.subtitles])
                    material.content = content
                    material.word_count = len(content)
                    material.success = True
                    return

            # 本地音频
            text = transcriber.transcribe(material.path)
            material.content = text
            material.word_count = len(text)
            material.language = self._detect_content_lang(text)
            material.success = True

        except Exception as e:
            material.error = f"音视频转录失败: {e}"

    def _detect_content_lang(self, content: str) -> ContentLanguage:
        """检测内容语言"""
        if not content:
            return ContentLanguage.UNKNOWN

        chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", content))
        total_chars = len(re.sub(r"\s", "", content))

        if total_chars == 0:
            return ContentLanguage.UNKNOWN

        chinese_ratio = chinese_chars / total_chars

        if chinese_ratio > 0.3:
            return ContentLanguage.ZH
        elif chinese_ratio > 0.1:
            return ContentLanguage.MIXED
        else:
            return ContentLanguage.EN
